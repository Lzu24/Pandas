
from docx.api import Document
import pandas as pd
from itertools import chain

document = Document(r'C:\file_path.docx')
tables = document.tables

def word_reader(tables): #load data
    big_data = []
    for table in document.tables:
        data = []
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
            
            big_data.append(data)

    ### cleaning data 
    df = pd.DataFrame(list(chain.from_iterable(big_data)))
    df = df.drop_duplicates()
    df = df.iloc[0:,:4]

    df["L.p."].fillna(df["Lp."], inplace=True)
    df = df.drop(df.index[df['L.p.'] == 'RAZEM'])
    df['L.p.'] = pd.to_numeric(df['L.p.'])

    df = df.rename(columns = {'Oryginał/\nrefabrykacja':'Oryginał/refabrykacja'})
    df = df[['L.p.','Opis przedmiotu zamówienia','Ilość sztuk','Oryginał/refabrykacja']]

    #file_name = 'word_frame.xls'
    #df.to_excel(file_name, header=True )
    return(df)
   

print(word_reader(tables))